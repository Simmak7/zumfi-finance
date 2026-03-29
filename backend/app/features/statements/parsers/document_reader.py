"""Unified document reader for PDF, Word, and image documents.

Provides a consistent interface for extracting text and tables from
PDFs, Word documents, and images (via Tesseract OCR).

For image-based PDFs (scanned statements, CSOB, etc.), automatically
falls back to OCR when pdfplumber extracts no text.
"""

import os
import logging
from typing import List, Optional

logger = logging.getLogger(__name__)

IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.webp', '.heic'}


class DocumentPage:
    """Represents a single page/section of a document."""

    def __init__(self, text: str, tables: Optional[List[List[List[str]]]] = None):
        self.text = text
        self._tables = tables or []

    def extract_text(self) -> str:
        """Extract text from this page."""
        return self.text

    def extract_tables(self) -> List[List[List[str]]]:
        """Extract tables from this page."""
        return self._tables


class Document:
    """Unified document interface for PDF, Word, and image files."""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.pages: List[DocumentPage] = []
        self._is_pdf = False
        self._load()

    def _load(self):
        """Load the document based on file extension."""
        _, ext = os.path.splitext(self.file_path.lower())

        if ext == '.pdf':
            self._is_pdf = True
            self._load_pdf()
        elif ext in ['.docx', '.doc']:
            self._is_pdf = False
            self._load_docx()
        elif ext in IMAGE_EXTENSIONS:
            self._load_image()
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _load_pdf(self):
        """Load PDF document using pdfplumber.

        Falls back to OCR (Tesseract) for pages with no extractable text,
        which handles scanned/image-based PDFs (CSOB, etc.).
        """
        import pdfplumber

        with pdfplumber.open(self.file_path) as pdf:
            # First pass: check if ANY page has extractable text chars
            has_any_text = any(page.chars for page in pdf.pages)

            for page in pdf.pages:
                text = page.extract_text() or ""
                tables = page.extract_tables() or []

                # If this page has no text and the PDF overall lacks text,
                # this is a scanned/image-based PDF — OCR it
                if not text.strip() and not has_any_text:
                    try:
                        pil_img = page.to_image(resolution=600).original
                        text = ocr_pdf_page(pil_img)
                        logger.info(
                            f"OCR fallback for PDF page: {len(text)} chars"
                        )
                    except Exception as e:
                        logger.warning(f"OCR fallback failed: {e}")

                self.pages.append(DocumentPage(text, tables))

    def _load_docx(self):
        """Load Word document using python-docx."""
        from docx import Document as DocxDocument

        doc = DocxDocument(self.file_path)

        # Extract paragraphs and tables
        text_parts = []
        extracted_tables = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract tables as list of lists (similar to pdfplumber format)
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                extracted_tables.append(table_data)

        # Treat entire document as one page
        full_text = '\n'.join(text_parts)

        # Create a single page with all content
        # (Word documents don't have the same page concept as PDFs)
        self.pages.append(DocumentPage(full_text, extracted_tables))

    def _load_image(self):
        """Load image and extract text via Tesseract OCR."""
        text = ocr_image(self.file_path)
        self.pages.append(DocumentPage(text, []))

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        pass


def ocr_pdf_page(pil_img) -> str:
    """OCR a single PDF page rendered as a PIL Image.

    Uses moderate pre-processing since PDF-rendered images are already
    clean (no camera distortion, good resolution). Contrast and
    sharpening improve recognition of small amounts and icons.
    """
    import pytesseract
    from PIL import ImageOps, ImageEnhance, ImageFilter

    # Convert to grayscale
    img = ImageOps.grayscale(pil_img)

    # Contrast boost + sharpen for better digit/symbol recognition
    img = ImageEnhance.Contrast(img).enhance(2.0)
    img = img.filter(ImageFilter.SHARPEN)

    return pytesseract.image_to_string(img, lang='ces+eng', config='--psm 6')


def ocr_image(file_path: str) -> str:
    """Extract text from an image using Tesseract OCR with pre-processing.

    Applies grayscale, upscaling, contrast enhancement, sharpening, and
    binarization to improve OCR accuracy on photos of bank statements.
    """
    import pytesseract
    from PIL import Image, ImageFilter, ImageEnhance, ImageOps

    img = Image.open(file_path)

    # Convert RGBA to RGB (remove alpha channel)
    if img.mode == 'RGBA':
        bg = Image.new('RGB', img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[3])
        img = bg

    # Convert to grayscale
    img = ImageOps.grayscale(img)

    # Upscale small images (Tesseract needs ~300 DPI for good results)
    w, h = img.size
    if w < 2000:
        scale = max(2, 2000 // w + 1)
        img = img.resize((w * scale, h * scale), Image.LANCZOS)

    # Increase contrast
    img = ImageEnhance.Contrast(img).enhance(2.0)

    # Sharpen
    img = img.filter(ImageFilter.SHARPEN)

    # Binarize (black & white threshold)
    img = img.point(lambda x: 0 if x < 140 else 255, '1')

    return pytesseract.image_to_string(img, lang='ces+eng', config='--psm 6')


_document_cache: dict[str, Document] = {}


def open_document(file_path: str) -> Document:
    """Open a document (PDF, Word, or image) and return a Document object.

    Uses a per-path cache so that detection + parsing don't OCR the same
    file twice. Cache is cleared after the document is retrieved once.

    Args:
        file_path: Path to the document file

    Returns:
        Document object with pages
    """
    real = os.path.realpath(file_path)
    if real in _document_cache:
        doc = _document_cache.pop(real)
        return doc
    return Document(file_path)


def preload_document(file_path: str) -> Document:
    """Pre-load and cache a document (used by detector to avoid double OCR)."""
    real = os.path.realpath(file_path)
    if real not in _document_cache:
        _document_cache[real] = Document(file_path)
    return _document_cache[real]

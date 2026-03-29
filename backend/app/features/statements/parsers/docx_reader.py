"""Word document (.docx) text extraction utility.

Provides functions to extract text from Word documents
similar to how pdfplumber extracts text from PDFs.
"""

from docx import Document


def extract_text_from_docx(docx_path: str) -> str:
    """Extract all text from a .docx file.

    Args:
        docx_path: Path to the .docx file

    Returns:
        Extracted text as a single string with newlines preserved
    """
    try:
        doc = Document(docx_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text_parts.append('\t'.join(row_text))

        return '\n'.join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract text from Word document: {e}")


def extract_first_page_text(docx_path: str, max_paragraphs: int = 20) -> str:
    """Extract text from the first 'page' (first N paragraphs) of a Word document.

    Args:
        docx_path: Path to the .docx file
        max_paragraphs: Number of paragraphs to consider as 'first page'

    Returns:
        Text from the first page
    """
    try:
        doc = Document(docx_path)
        text_parts = []

        paragraph_count = 0
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
                paragraph_count += 1
                if paragraph_count >= max_paragraphs:
                    break

        return '\n'.join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract text from Word document: {e}")
